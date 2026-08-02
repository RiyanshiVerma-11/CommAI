import sqlite3
import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets cell shading background color."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_borders(cell):
    """Adds thin grey borders to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows):
    """Creates a beautifully styled report table in the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_borders(hdr_cells[i])
        # Text formatting
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Segoe UI"
        run.font.size = Pt(9.5)

    # Populate rows with zebra shading
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        cells = row.cells
        bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val or "")
            set_cell_background(cells[c_idx], bg_color)
            set_cell_borders(cells[c_idx])
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = "Segoe UI"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 51, 51)
                
    doc.add_paragraph() # Spacing

def export_word(db_path: str, output_path: str):
    """
    Connects to the SQLite database and exports an audit and compliance report in Word format (.docx).
    """
    if not os.path.exists(db_path):
        print(f"Error: Database file does not exist at '{db_path}'")
        return False

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    doc = Document()
    
    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CommAI mass communication platform")
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.name = "Segoe UI"
    title_run.font.color.rgb = RGBColor(30, 58, 138)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Operational Database Audit & Compliance Report")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.name = "Segoe UI"
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f"Export Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Database File: {os.path.basename(db_path)}\n"
        f"Authority: CommAI System Administrator\n"
    )
    meta_run.font.name = "Segoe UI"
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 1: EXEC SUMMARY & KEY STATS
    # ----------------------------------------------------
    h1 = doc.add_heading("1. Executive summary statistics", level=1)
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h1.runs[0].font.name = "Segoe UI"
    
    # Query stats
    def get_count(table_name):
        try:
            cursor.execute(f"SELECT COUNT(*) FROM {table_name};")
            return cursor.fetchone()[0]
        except Exception:
            return 0

    total_users = get_count("users")
    total_audiences = get_count("audiences")
    total_campaigns = get_count("campaigns")
    total_logs = get_count("delivery_logs")
    total_sos = get_count("sos_reports")
    
    summary_p = doc.add_paragraph()
    summary_p.add_run(
        f"This report outlines the operational metrics retrieved from the CommAI SQLite database repository. "
        f"CommAI enables multi-channel emergency broadcasting and awareness campaigns in local languages. "
        f"Below are the high-level platform status counters:\n"
    ).font.name = "Segoe UI"
    
    stats_headers = ["Metric Parameter", "Count Indicator"]
    stats_rows = [
        ["Total Platform Operators", total_users],
        ["Registered Audience Citizens", total_audiences],
        ["Broadcasting Campaigns Created", total_campaigns],
        ["Message Delivery Log Records", total_logs],
        ["SOS Hazards reported", total_sos]
    ]
    add_styled_table(doc, stats_headers, stats_rows)

    # ----------------------------------------------------
    # SECTION 2: CAMPAIGN LOGS
    # ----------------------------------------------------
    h2 = doc.add_heading("2. Recent campaigns overview", level=1)
    h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h2.runs[0].font.name = "Segoe UI"
    
    try:
        cursor.execute("SELECT id, title, campaign_type, status, target_audience_count, created_at FROM campaigns ORDER BY created_at DESC LIMIT 10;")
        camps = cursor.fetchall()
    except Exception:
        camps = []
        
    camp_headers = ["ID", "Title", "Type", "Status", "Target Count", "Created Date"]
    add_styled_table(doc, camp_headers, camps)

    # ----------------------------------------------------
    # SECTION 3: PUBLIC SOS HAZARDS
    # ----------------------------------------------------
    h3 = doc.add_heading("3. SOS incident queue triage log", level=1)
    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h3.runs[0].font.name = "Segoe UI"
    
    try:
        cursor.execute("SELECT id, title, report_type, status, location_name, created_at FROM sos_reports ORDER BY created_at DESC LIMIT 15;")
        sos_list = cursor.fetchall()
    except Exception:
        sos_list = []
        
    sos_headers = ["ID", "Title / Incident", "Type", "Triage Status", "Location Reference", "Timestamp"]
    add_styled_table(doc, sos_headers, sos_list)

    # Save document
    doc.save(output_path)
    conn.close()
    print(f"Word audit report successfully exported to '{output_path}'")
    return True

if __name__ == "__main__":
    # If run standalone, detect DB path
    base_dir = os.path.dirname(os.path.abspath(__file__))
    db = os.path.join(base_dir, "comm_platform.db")
    out = os.path.join(base_dir, "database_export.docx")
    export_word(db, out)
