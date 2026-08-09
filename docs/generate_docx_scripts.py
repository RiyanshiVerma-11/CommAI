import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def convert_md_to_docx(md_path, docx_path, title_text):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return

    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(33, 37, 41)

    # Document Title Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(title_text)
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(14, 116, 144) # Cyan/Teal accent

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_blockquote = False
    blockquote_text = []

    for raw_line in lines:
        line = raw_line.rstrip()

        # Handle Blockquotes (> )
        if line.startswith('> '):
            in_blockquote = True
            text_chunk = line[2:].strip()
            blockquote_text.append(text_chunk)
            continue
        else:
            if in_blockquote:
                # Flush accumulated blockquote
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                cell = table.cell(0, 0)
                set_cell_background(cell, "F0F9FF") # Light Blue background
                cell.width = Inches(6.8)
                
                bq_p = cell.paragraphs[0]
                bq_p.paragraph_format.space_before = Pt(6)
                bq_p.paragraph_format.space_after = Pt(6)
                bq_p.paragraph_format.left_indent = Inches(0.1)
                
                bq_run = bq_p.add_run("\n".join(blockquote_text))
                bq_run.font.name = 'Georgia'
                bq_run.font.size = Pt(11)
                bq_run.font.italic = True
                bq_run.font.color.rgb = RGBColor(30, 58, 138) # Dark blue text
                
                doc.add_paragraph() # Spacing after blockquote
                in_blockquote = False
                blockquote_text = []

        if not line.strip():
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:].strip())
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:].strip())
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Strip markdown formatting stars
            text = line[2:].strip()
            # Simple bold handler
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            text = re.sub(r'^\d+\.\s', '', line).strip()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Generated DOCX: {docx_path}")

if __name__ == "__main__":
    docs_dir = r"d:\Riyanshi\01_coding\projects\32 CommAI\docs"
    
    teleprompter_md = os.path.join(docs_dir, "LIVE_DEMO_TELEPROMPTER_SCRIPT.md")
    teleprompter_docx = os.path.join(docs_dir, "LIVE_DEMO_TELEPROMPTER_SCRIPT.docx")
    convert_md_to_docx(teleprompter_md, teleprompter_docx, "CommAI - Live Video & Screen-Share Teleprompter Script")

    master_md = os.path.join(docs_dir, "COMM_AI_MASTER_DEMO_SCRIPT.md")
    master_docx = os.path.join(docs_dir, "COMM_AI_MASTER_DEMO_SCRIPT.docx")
    convert_md_to_docx(master_md, master_docx, "CommAI - Master 1-Hour Presentation & Technical Guide")
